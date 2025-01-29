import sys
import os
from dotenv import load_dotenv
from crewai import Agent, Task, Crew
from langchain_openai import ChatOpenAI
from docx import Document
import json
from typing import List, Dict, Any
import logging
from datetime import datetime

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('test_generation.log'),
        logging.StreamHandler()
    ]
)

# Load environment variables
load_dotenv()

OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY is missing. Ensure it is set in the environment or .env file.")

class TestCaseGenerationCrew:
    def __init__(self, requirements_text: str):
        """Initialize the Test Case Generation Crew using GPT as the LLM."""
        self.requirements_text = requirements_text
        self.llm = ChatOpenAI(
            api_key=OPENAI_API_KEY,
            model="gpt-4-turbo",
            temperature=0.1
        )
        self.output_dir = self._create_output_directory()
        logging.info("TestCaseGenerationCrew initialized")

    def _create_output_directory(self) -> str:
        """Create a timestamped output directory for artifacts."""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_dir = f"test_artifacts_{timestamp}"
        os.makedirs(output_dir, exist_ok=True)
        return output_dir

    def create_requirements_analyzer(self) -> Agent:
        """Create the Requirements Analyzer agent."""
        return Agent(
            role='Requirements Analyst',
            goal='Analyze software requirements and extract structured information for test case creation',
            backstory='''You are a senior Requirements Analyst with 15+ years of experience in 
            analyzing complex software requirements. You excel at identifying edge cases, 
            non-functional requirements, and hidden constraints.''',
            verbose=True,
            llm=self.llm
        )

    def create_test_case_generator(self) -> Agent:
        """Create the Test Case Generator agent."""
        return Agent(
            role='Test Case Engineer',
            goal='Generate comprehensive test cases covering all requirement aspects',
            backstory='''You are a senior Test Case Engineer specializing in creating 
            detailed test cases for complex systems. Your test cases are known for their 
            clarity, completeness, and attention to edge cases.''',
            verbose=True,
            llm=self.llm
        )

    def create_test_data_generator(self) -> Agent:
        """Create the Test Data Generator agent."""
        return Agent(
            role='Test Data Engineer',
            goal='Generate comprehensive test data sets for all test scenarios',
            backstory='''You are a specialized Test Data Engineer with expertise in creating 
            realistic and comprehensive test data. You understand data patterns, boundary 
            conditions, and security implications of test data.''',
            verbose=True,
            llm=self.llm
        )

    def create_test_case_validator(self) -> Agent:
        """Create the Test Case Validator agent."""
        return Agent(
            role='Test Case Validator',
            goal='Validate and optimize test cases for completeness and effectiveness',
            backstory='''You are a Test Case Validation expert who ensures test cases meet 
            quality standards and provide optimal coverage. You excel at identifying gaps 
            and optimizing test suites.''',
            verbose=True,
            llm=self.llm
        )

    def analyze_requirements_task(self) -> Task:
        """Create the requirements analysis task."""
        output_file = os.path.join(self.output_dir, 'requirements_analysis.md')
        return Task(
            description=f"""Thoroughly analyze the following software requirements document:
            {self.requirements_text}

            Extract and structure:
            1. Functional Requirements
            2. Non-functional Requirements
            3. Business Rules and Constraints
            4. User Scenarios and Workflows
            5. System Interfaces and Dependencies
            6. Performance Requirements
            7. Security Requirements
            8. Edge Cases and Boundary Conditions

            Provide a detailed, structured analysis suitable for test case creation.""",
            agent=self.create_requirements_analyzer(),
            output_file=output_file
        )

    def generate_test_cases_task(self) -> Task:
        """Create the test case generation task."""
        output_file = os.path.join(self.output_dir, 'detailed_test_cases.md')
        return Task(
            description="""Generate comprehensive test cases based on the requirements analysis.
            
            Include test cases for:
            1. Functional Requirements
            2. Non-functional Requirements
            3. Edge Cases and Boundary Conditions
            4. Error Scenarios
            5. Integration Points
            6. Security Scenarios
            7. Performance Testing

            Each test case must include:
            - Unique ID and Title
            - Requirement Reference
            - Preconditions
            - Test Steps
            - Expected Results
            - Test Data References
            - Priority and Type
            - Notes and Risks""",
            agent=self.create_test_case_generator(),
            dependencies=[os.path.join(self.output_dir, 'requirements_analysis.md')],
            output_file=output_file
        )

    def generate_test_data_task(self) -> Task:
        """Create the test data generation task."""
        output_file = os.path.join(self.output_dir, 'test_data_sets.json')
        return Task(
            description="""Generate comprehensive test data sets for all test cases.
            
            Include data for:
            1. Happy Path Scenarios
            2. Boundary Values
            3. Invalid Inputs
            4. Edge Cases
            5. Security Testing
            6. Performance Testing

            Provide data sets in structured JSON format with:
            - Test Case ID Reference
            - Input Data Sets
            - Expected Results
            - Data Dependencies
            - Environmental Requirements""",
            agent=self.create_test_data_generator(),
            dependencies=[os.path.join(self.output_dir, 'detailed_test_cases.md')],
            output_file=output_file
        )

    def validate_test_cases_task(self) -> Task:
        """Create the test case validation task."""
        output_file = os.path.join(self.output_dir, 'validation_report.md')
        return Task(
            description="""Validate and optimize the test suite.
            
            Perform:
            1. Requirements Coverage Analysis
            2. Test Case Quality Assessment
            3. Test Suite Optimization
            4. Risk Analysis
            
            Provide detailed report including:
            - Coverage Metrics
            - Quality Metrics
            - Optimization Recommendations
            - Risk Assessment
            - Improvement Suggestions""",
            agent=self.create_test_case_validator(),
            dependencies=[
                os.path.join(self.output_dir, 'detailed_test_cases.md'),
                os.path.join(self.output_dir, 'test_data_sets.json')
            ],
            output_file=output_file
        )

    def generate_test_cases(self) -> Dict[str, Any]:
        """Execute the complete test case generation workflow."""
        try:
            logging.info("Starting test case generation workflow")
            
            # Create tasks
            tasks = [
                self.analyze_requirements_task(),
                self.generate_test_cases_task(),
                self.generate_test_data_task(),
                self.validate_test_cases_task()
            ]

            # Create agents
            agents = [
                self.create_requirements_analyzer(),
                self.create_test_case_generator(),
                self.create_test_data_generator(),
                self.create_test_case_validator()
            ]

            # Create and execute crew
            crew = Crew(
                agents=agents,
                tasks=tasks,
                verbose=True
            )

            # Execute workflow
            result = crew.kickoff()
            
            # Export results
            self.export_results(result)
            
            logging.info("Test case generation workflow completed successfully")
            
            return {
                'status': 'success',
                'result': result,
                'artifacts_directory': self.output_dir
            }
            
        except Exception as e:
            logging.error(f"Error in test case generation: {str(e)}")
            return {
                'status': 'error',
                'error': str(e)
            }

    def export_results(self, results: Any) -> None:
        """Export results to various formats."""
        try:
            # Export to Word
            self.export_to_word(results)
            
            # Export to JSON
            self.export_to_json(results)
            
            logging.info("Results exported successfully")
            
        except Exception as e:
            logging.error(f"Error exporting results: {str(e)}")
            raise

    def export_to_word(self, results: Any) -> str:
        """Export results to a Word document."""
        doc = Document()
        output_file = os.path.join(self.output_dir, 'Generated_Test_Cases.docx')

        # Add title
        doc.add_heading('Test Case Generation Results', level=1)

        # Add requirements analysis
        doc.add_heading('Requirements Analysis', level=2)
        doc.add_paragraph(str(results.get('requirements_analysis', '')))

        # Add test cases
        doc.add_heading('Test Cases', level=2)
        doc.add_paragraph(str(results.get('test_cases', '')))

        # Add test data
        doc.add_heading('Test Data', level=2)
        doc.add_paragraph(str(results.get('test_data', '')))

        # Add validation report
        doc.add_heading('Validation Report', level=2)
        doc.add_paragraph(str(results.get('validation_report', '')))

        doc.save(output_file)
        return output_file

    def export_to_json(self, results: Any) -> str:
        """Export results to JSON format."""
        output_file = os.path.join(self.output_dir, 'test_cases_complete.json')
        
        with open(output_file, 'w') as f:
            json.dump(results, f, indent=4)
        
        return output_file

def main(file_path: str) -> int:
    """Main function to execute the test case generation workflow."""
    try:
        # Read the input document
        doc = Document(file_path)
        requirements_content = "\n".join([para.text for para in doc.paragraphs])

        # Initialize the crew
        crew = TestCaseGenerationCrew(requirements_content)
        
        # Generate test cases
        result = crew.generate_test_cases()
        
        # Print results
        print(json.dumps(result))
        
        return 0 if result['status'] == 'success' else 1
        
    except Exception as e:
        print(json.dumps({'error': str(e)}), file=sys.stderr)
        return 1

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python test_case_generator.py <requirements_file_path>")
        sys.exit(1)
    sys.exit(main(sys.argv[1]))